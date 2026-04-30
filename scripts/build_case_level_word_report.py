from pathlib import Path
import pandas as pd
from PIL import Image, ImageOps, ImageDraw, ImageFont

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Reconstructing_Wonderland_Case_Level_Report.docx"
DERIVED = ROOT / "derived_report_figures"


def pth(*parts):
    return ROOT.joinpath(*parts)


FIG = ROOT / "outputs" / "figures"
GEN = ROOT / "report_package" / "generated_analysis"
STILLS = ROOT / "report_package"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=90, start=90, bottom=90, end=90):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D0D0D0", size="6"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_run(run, size=None, bold=None, italic=None, color=None):
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")


def add_para(doc, text="", style=None, space_after=4, first_line=False):
    para = doc.add_paragraph(style=style)
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.line_spacing = 1.05
    if first_line:
        para.paragraph_format.first_line_indent = Cm(0.45)
    if text:
        run = para.add_run(text)
        style_run(run, size=10.5)
    return para


def add_heading(doc, text, level=1):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(8 if level == 1 else 5)
    para.paragraph_format.space_after = Pt(5 if level == 1 else 3)
    run = para.add_run(text)
    if level == 1:
        style_run(run, size=16, bold=True, italic=True, color="111111")
    elif level == 2:
        style_run(run, size=12.5, bold=True, italic=True, color="111111")
    else:
        style_run(run, size=11, bold=True, italic=True, color="111111")
    return para


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color="B7C0C7", size="8")
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F3F6F8")
    set_cell_margins(cell, top=140, start=160, bottom=140, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    style_run(r, size=10.5, bold=True, italic=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.line_spacing = 1.05
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    style_run(r2, size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table_from_rows(doc, headers, rows, widths=None, caption=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_borders(table)
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_shading(hdr[i], "E9EEF2")
        set_cell_margins(hdr[i])
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        r = hdr[i].paragraphs[0].add_run(str(h))
        style_run(r, size=8.8, bold=True)
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_repeat_table_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.line_spacing = 1.0
            r = p.add_run(str(val))
            style_run(r, size=8.6)
            if i > 0 and len(str(val)) < 18:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    if caption:
        add_caption(doc, caption)
    return table


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    style_run(r, size=8.5, italic=True, color="333333")
    return p


def add_image(doc, path, caption, width_cm=15.5):
    path = Path(path)
    if not path.exists():
        add_callout(doc, "Missing figure placeholder", f"Expected image: {path}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    add_caption(doc, caption)


def make_collage(left_path, left_caption, right_path, right_caption, name):
    DERIVED.mkdir(parents=True, exist_ok=True)
    out = DERIVED / name
    items = [(Path(left_path), left_caption), (Path(right_path), right_caption)]
    panels = []
    for path, caption in items:
        if path.exists():
            im = Image.open(path).convert("RGB")
        else:
            im = Image.new("RGB", (1200, 700), "white")
            d = ImageDraw.Draw(im)
            d.text((40, 40), f"Missing: {path}", fill=(0, 0, 0))
        im = ImageOps.contain(im, (1000, 520), method=Image.Resampling.LANCZOS)
        panel = Image.new("RGB", (1060, 650), "white")
        panel.paste(im, ((1060 - im.width) // 2, 20))
        d = ImageDraw.Draw(panel)
        d.text((30, 565), caption, fill=(40, 40, 40))
        panels.append(panel)
    collage = Image.new("RGB", (2120, 650), "white")
    collage.paste(panels[0], (0, 0))
    collage.paste(panels[1], (1060, 0))
    collage.save(out)
    return out


def add_two_images(doc, left_path, left_caption, right_path, right_caption):
    safe_name = f"collage_{abs(hash((str(left_path), str(right_path), left_caption, right_caption))) % 10**12}.png"
    collage = make_collage(left_path, left_caption, right_path, right_caption, safe_name)
    add_image(doc, collage, f"{left_caption} / {right_caption}", width_cm=15.5)


def add_static_toc(doc):
    add_heading(doc, "Report Structure", 2)
    items = [
        "1. Project overview and Wonderland design logic",
        "2. Part 1: dataset construction and cleaning",
        "3. Text vectorisation: motif extraction and semantic retrieval",
        "4. Image and audio vectorisation: atmosphere and motion",
        "5. Cross-modal model comparison and design decision",
        "6. Extended Part 1 figure analysis",
        "7. Parameter mapping: from analysis to design control",
        "8. Workflow 1: Blender fragment library",
        "9. Workflow 2: TouchDesigner procedural system",
        "10. Final animation and critical reflection",
        "11. Appendix: evidence checklist and credits",
    ]
    for item in items:
        p = add_para(doc)
        r = p.add_run(item)
        style_run(r, size=10.3)


def add_part1_extended_figure_analysis(doc):
    add_page_break(doc)
    add_heading(doc, "6. Extended Part 1 Figure Analysis", 1)
    add_para(
        doc,
        "This section expands Part 1 by reading the notebook figures as evidence rather than as decoration. The figures are grouped by workflow stage: data collection, data processing, vectorisation and model comparison. The purpose is to show how each plotted output contributes to the Alice/Wonderland design logic.",
    )
    add_callout(
        doc,
        "How to read this section",
        "A strong data-driven design report should not simply paste charts. Each figure needs a reading: what it proves about the dataset, what limitation it reveals, and how it changes the design workflow.",
    )

    add_heading(doc, "6.1 Data collection figures: building the Wonderland evidence base", 2)
    add_para(
        doc,
        "The data-collection figures establish the scale and balance of the source material. The dataset-size overview proves that all three modalities are large enough for the assignment requirement. The text word-count and chapter-distribution figures show whether the narrative fragments are short enough for vectorisation while still distributed across the Wonderland story. The image and audio figures prove that the visual and temporal datasets are not isolated additions, but parallel sources for atmosphere and motion.",
    )
    add_two_images(
        doc,
        FIG / "part1_data_collection" / "07_dataset_size_overview.png",
        "Dataset size overview from the collection notebook.",
        FIG / "part1_data_collection" / "01_text_word_count_distribution.png",
        "Raw text word-count distribution.",
    )
    add_para(
        doc,
        "The word-count figure is important because it checks whether fragments are suitable for semantic embedding. Very long passages would mix several motifs at once, while extremely short fragments could become too sparse for meaningful retrieval. For an Alice/Wonderland project, fragment length affects whether a passage represents a specific motif such as a locked door, the White Rabbit, the Queen's command or the Mad Tea-Party, rather than a blurred multi-scene summary.",
    )
    add_two_images(
        doc,
        FIG / "part1_data_collection" / "02_text_segments_per_chapter.png",
        "Raw text segments per chapter.",
        FIG / "part1_data_collection" / "03_image_category_distribution.png",
        "Raw image category distribution.",
    )
    add_para(
        doc,
        "The chapter-distribution figure helps check whether the text dataset overrepresents only the opening rabbit-hole sequence. A strong Wonderland reconstruction needs evidence from multiple narrative states: descent, scale change, animal encounters, tea-time, Queen/court authority and the waking dream frame. The image category distribution provides a parallel check for visual balance. If one category dominated, the generated material palette would risk becoming visually one-note.",
    )
    add_two_images(
        doc,
        FIG / "part1_data_collection" / "04_image_sample_grid_with_filenames.png",
        "Raw image sample grid with filenames.",
        FIG / "part1_data_collection" / "04_curated_image_sample_grid.png",
        "Curated image sample grid.",
    )
    add_para(
        doc,
        "The image grids are more than a visual appendix. They allow the reader to verify that the dataset contains a recognisable Wonderland atmosphere: character references, surreal objects, spatial thresholds and dreamlike scenes. The filename version is useful for traceability, while the curated grid is useful for communicating design atmosphere.",
    )
    add_two_images(
        doc,
        FIG / "part1_data_collection" / "05_audio_category_distribution.png",
        "Raw audio category distribution.",
        FIG / "part1_data_collection" / "06_audio_duration_distribution_clipped.png",
        "Clipped audio duration distribution.",
    )
    add_para(
        doc,
        "The audio category and duration plots show whether sound material can support motion design. Duration matters because extremely long clips would dominate temporal mapping, while very short clips may produce abrupt behaviour. For Wonderland, both are conceptually relevant: sudden changes can express instability, but the final animation still needs controlled rhythm.",
    )

    add_heading(doc, "6.2 Data processing figures: proving that the material is usable", 2)
    add_para(
        doc,
        "The processing figures repeat some collection checks after cleaning. This repetition is useful because it proves that cleaning did not damage the dataset structure. In a data-driven design workflow, cleaning must reduce noise while preserving the narrative and visual range needed for design.",
    )
    add_two_images(
        doc,
        FIG / "part1_data_processing" / "01_cleaned_text_word_count_distribution.png",
        "Cleaned text word-count distribution.",
        FIG / "part1_data_processing" / "02_cleaned_text_segments_per_chapter.png",
        "Cleaned text segments per chapter.",
    )
    add_para(
        doc,
        "The cleaned text plots confirm that the narrative structure remains intact after preprocessing. This is essential for later retrieval analysis. The report can argue that semantic results are based on cleaned, validated fragments rather than arbitrary raw text.",
    )
    add_two_images(
        doc,
        FIG / "part1_data_processing" / "03_cleaned_image_category_distribution.png",
        "Cleaned image category distribution.",
        FIG / "part1_data_processing" / "04_cleaned_image_size_distribution.png",
        "Cleaned image size distribution.",
    )
    add_two_images(
        doc,
        FIG / "part1_data_processing" / "05_cleaned_image_aspect_ratio_distribution.png",
        "Cleaned image aspect-ratio distribution.",
        FIG / "part1_data_processing" / "06_curated_wonderland_image_showcase_with_filenames.png",
        "Curated Wonderland image showcase with filenames.",
    )
    add_para(
        doc,
        "The image size and aspect-ratio figures matter because they affect feature extraction. If image dimensions or aspect ratios were extremely inconsistent, handcrafted visual descriptors could partly reflect file format rather than atmosphere. The curated showcase with filenames supports both theme communication and traceability: the reader can see the Wonderland visual field and still locate source files.",
    )
    add_two_images(
        doc,
        FIG / "part1_data_processing" / "07_cleaned_audio_category_distribution.png",
        "Cleaned audio category distribution.",
        FIG / "part1_data_processing" / "08_cleaned_audio_duration_distribution.png",
        "Cleaned audio duration distribution.",
    )
    add_para(
        doc,
        "The cleaned audio plots support the later decision to map sound into behaviour. The distribution of durations and categories indicates whether there is enough temporal variation to produce changes in speed, amplitude and sharpness. This directly prepares the TouchDesigner workflow, where audio-derived parameters control procedural movement.",
    )
    add_image(
        doc,
        FIG / "part1_data_processing" / "09_raw_vs_cleaned_dataset_sizes.png",
        "Raw versus cleaned dataset sizes. This figure summarises retention and validates the preprocessing stage.",
    )
    add_para(
        doc,
        "The raw-versus-cleaned figure is a compact quality-control proof. Text and images remain stable, while audio loses one invalid sample. This small removal is useful to mention because it shows that the pipeline checks data quality before design translation.",
    )

    add_heading(doc, "6.3 Text vectorisation figures: motif extraction and semantic ambiguity", 2)
    add_para(
        doc,
        "The text vectorisation figures should be read through the Alice theme. TF-IDF identifies explicit Wonderland motifs, while embedding plots and heatmaps reveal how difficult it is to separate dreamlike narrative fragments into clean clusters.",
    )
    add_two_images(
        doc,
        FIG / "part1_vectorisation" / "00_cleaned_dataset_sizes.png",
        "Cleaned dataset sizes used before vectorisation.",
        FIG / "part1_vectorisation" / "01_dataset_sizes.png",
        "Vectorisation dataset sizes.",
    )
    add_two_images(
        doc,
        FIG / "part1_vectorisation" / "01_text_top_tfidf_terms.png",
        "Top TF-IDF terms for the Alice text dataset.",
        FIG / "part1_vectorisation" / "03_text_pca_explained_variance.png",
        "Text PCA explained variance.",
    )
    add_para(
        doc,
        "The top TF-IDF terms work almost like a motif index: Alice, Queen, Time, Hatter, Rabbit, Door, Cat and Caterpillar. These are not generic words; they are symbolic anchors for Wonderland. The PCA explained variance figure shows that text structure is distributed across many dimensions, which is expected for a literary dataset where motifs overlap.",
    )
    add_two_images(
        doc,
        FIG / "part1_vectorisation" / "02_text_vectorisation_pca_comparison.png",
        "TF-IDF versus MiniLM PCA comparison.",
        FIG / "part1_vectorisation" / "02_text_cluster_comparison_heatmap.png",
        "Text cluster comparison heatmap.",
    )
    add_two_images(
        doc,
        FIG / "part1_vectorisation" / "03_text_embedding_cluster_heatmap.png",
        "Text embedding cluster heatmap.",
        GEN / "text_retrieval_top_results.png",
        "Top retrieval examples for Wonderland motifs.",
    )
    add_para(
        doc,
        "The heatmaps and retrieval figure should be interpreted together. The clusters are not sharply separated, but semantic retrieval still improves the design workflow. This is the key case-level point: the project should not force Wonderland into rigid categories. Instead, it should use semantic similarity to create soft scene tendencies such as curious, dreamlike, anxious and chaotic.",
    )

    add_heading(doc, "6.4 Image vectorisation figures: from visual descriptors to material atmosphere", 2)
    add_para(
        doc,
        "The image vectorisation figures explain how the visual dataset becomes material logic. Handcrafted features cannot recognise symbolic objects on their own, but they can describe qualities that matter for design: brightness, saturation, edge density, size and visual complexity.",
    )
    add_two_images(
        doc,
        FIG / "part1_vectorisation" / "04_image_handcrafted_pca_projection.png",
        "Handcrafted image PCA projection.",
        FIG / "part1_vectorisation" / "05_image_feature_distributions.png",
        "Image feature distributions.",
    )
    add_two_images(
        doc,
        FIG / "part1_vectorisation" / "04_image_feature_correlation.png",
        "Image feature correlation.",
        FIG / "part1_vectorisation" / "06_image_cluster_heatmap.png",
        "Image cluster heatmap.",
    )
    add_para(
        doc,
        "The PCA projection and cluster heatmap show whether images form visually coherent groups. The correlation and distribution plots explain which visual properties vary independently. This matters because independent variation creates a richer material library. For example, if brightness and saturation were always locked together, the material system would be less flexible. The correlation figure shows enough separation to justify mapping visual features into roughness, colour intensity and texture choice.",
    )

    add_heading(doc, "6.5 Audio vectorisation figures: sound as animation behaviour", 2)
    add_para(
        doc,
        "The audio figures explain why broad descriptors are selected for motion mapping. In a Wonderland animation, sound is not only background atmosphere. It becomes the logic of instability, urgency and temporal distortion.",
    )
    add_two_images(
        doc,
        FIG / "part1_vectorisation" / "06_audio_basic_feature_correlation.png",
        "Basic audio feature correlation.",
        FIG / "part1_vectorisation" / "07_audio_basic_feature_correlation.png",
        "Basic audio feature correlation, alternate view.",
    )
    add_two_images(
        doc,
        FIG / "part1_vectorisation" / "08_audio_vectorisation_pca_comparison.png",
        "Audio vectorisation PCA comparison.",
        FIG / "part1_vectorisation" / "08_audio_pca_explained_variance.png",
        "Audio PCA explained variance.",
    )
    add_two_images(
        doc,
        FIG / "part1_vectorisation" / "07_audio_mfcc_cluster_heatmap.png",
        "MFCC cluster heatmap.",
        FIG / "part1_vectorisation" / "09_audio_mfcc_cluster_heatmap.png",
        "MFCC cluster heatmap, alternate view.",
    )
    add_para(
        doc,
        "The audio PCA and heatmap figures support the numerical comparison: basic descriptors are more coherent and easier to interpret as motion controls, while MFCCs describe timbral texture but are less direct for animation. This is why basic descriptors are mapped to motion_amp, anim_speed and temporal_sharpness.",
    )
    add_image(
        doc,
        FIG / "part1_vectorisation" / "10_feature_dimension_comparison.png",
        "Feature dimension comparison across modalities.",
    )
    add_para(
        doc,
        "The feature-dimension figure helps prevent a simplistic reading of the models. Higher dimensionality is not automatically better. MiniLM has more dimensions than TF-IDF and performs better semantically, but image and audio features are lower-dimensional and still more directly useful for design translation. The report should therefore evaluate representations through design usefulness, not only vector size.",
    )

    add_heading(doc, "6.6 Part 1 figure inventory", 2)
    inventory = [
        ["Collection: text word count", "Checks fragment length before embedding; supports readable motif-level text units."],
        ["Collection: text segments per chapter", "Checks narrative coverage across Wonderland chapters."],
        ["Collection: image category distribution", "Checks whether visual references are balanced."],
        ["Collection: image grids", "Communicates visual atmosphere and verifies source traceability."],
        ["Collection: audio category / duration", "Checks temporal material for motion mapping."],
        ["Processing: cleaned text plots", "Shows cleaning preserved narrative structure."],
        ["Processing: cleaned image size/aspect", "Prevents file-format bias in visual descriptors."],
        ["Processing: curated image showcase", "Connects dataset evidence to Wonderland aesthetics."],
        ["Processing: cleaned audio plots", "Supports later mapping from sound to behaviour."],
        ["Vectorisation: TF-IDF terms", "Explicit motif layer for Alice, Queen, Rabbit, Time, Door and other symbols."],
        ["Vectorisation: text PCA / heatmaps", "Shows semantic overlap and justifies soft scene mapping."],
        ["Vectorisation: image PCA / correlation", "Links visual descriptors to material atmosphere."],
        ["Vectorisation: audio PCA / heatmaps", "Justifies basic audio descriptors as motion controls."],
        ["Vectorisation: feature dimensions", "Shows why representation choice must consider design role, not only size."],
    ]
    add_table_from_rows(
        doc,
        ["Figure group", "Detailed report use"],
        inventory,
        widths=[5.2, 10.8],
        caption="Table 8. Part 1 figure inventory and analytical use.",
    )


def add_page_break(doc):
    doc.add_page_break()


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

    return doc


def build():
    doc = setup_document()

    # Cover
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(80)
    r = p.add_run("Reconstructing Wonderland")
    style_run(r, size=26, bold=True, italic=True)
    p2 = doc.add_paragraph()
    r2 = p2.add_run("A Data-Driven Spatial Narrative from Alice's Adventures in Wonderland")
    style_run(r2, size=15, bold=True, italic=True, color="333333")
    add_para(doc, "")
    add_callout(
        doc,
        "Report position",
        "This report documents how text, image and audio datasets are transformed into spatial fragments, material logic and animated behaviour. The focus is Alice/Wonderland as a system of falling, doors, time anxiety, dream logic, authority and scale change.",
    )
    add_image(
        doc,
        STILLS / "09_final_animation_reflection" / "stills" / "final_still_01.png",
        "Final animation still used as a visual anchor for the project outcome.",
        width_cm=14.5,
    )
    add_static_toc(doc)
    add_page_break(doc)

    # Chapter 1
    add_heading(doc, "1. Project Overview and Wonderland Design Logic", 1)
    add_para(
        doc,
        "The project explores how the world of Alice's Adventures in Wonderland can be reconstructed through data-driven design. The intention is not to illustrate the story literally, but to convert its narrative, visual and temporal qualities into a spatial system. Wonderland is defined by repeated transformations: Alice falls through the rabbit hole, changes scale, moves through locked doors, encounters unstable rules, and enters scenes where time, authority and identity behave strangely.",
    )
    add_para(
        doc,
        "These motifs make the story suitable for a multimodal workflow. Text captures narrative meaning and symbolic motifs. Images capture atmosphere, colour and material qualities. Audio captures rhythm, urgency and instability. The final design workflow therefore treats data as design evidence rather than as decoration.",
    )
    add_callout(
        doc,
        "Core argument",
        "Text controls meaning, image controls appearance, and audio controls behaviour. The parameter table is the bridge that turns those analyses into Blender and TouchDesigner design instructions.",
    )
    add_image(
        doc,
        GEN / "data_to_fragment_mapping_diagram.png",
        "Data-to-fragment workflow. The project moves from multimodal datasets to cleaned features, comparison, parameter mapping and design production.",
    )
    add_para(
        doc,
        "This report is structured around that chain. Part 1 establishes the dataset and model evidence. Part 2 explains how the evidence is translated into Blender, TouchDesigner and final animation workflows. This order matters because it lets the examiner follow the logic from raw data to spatial output.",
    )
    add_page_break(doc)

    # Chapter 2
    add_heading(doc, "2. Part 1: Dataset Construction and Cleaning", 1)
    add_heading(doc, "2.1 Dataset roles", 2)
    add_para(
        doc,
        "Three datasets were prepared for the project. The text dataset contains 288 cleaned narrative records and provides the semantic foundation for scene labels, emotional categories and motif retrieval. The image dataset contains 298 cleaned image records in the data-processing summary, with 278 records used in the vectorisation stage after feature-level filtering. The audio dataset contains 270 cleaned audio records after one raw sample was removed during validation.",
    )
    add_table_from_rows(
        doc,
        ["Dataset", "Cleaned / vectorised count", "Wonderland role", "Design translation"],
        [
            ["Text", "288", "Narrative motifs: rabbit-hole, door, time, Queen, dream, tea-party", "Emotion, scene label, prompt logic, fragment family"],
            ["Images", "298 cleaned / 278 vectorised", "Visual atmosphere and material cues", "Colour intensity, roughness, texture and visual density"],
            ["Audio", "270", "Rhythm, urgency and instability", "Motion amplitude, speed and temporal sharpness"],
        ],
        widths=[2.4, 3.2, 5.2, 5.2],
        caption="Table 1. Dataset scale and design role.",
    )
    add_image(
        doc,
        FIG / "part1_data_processing" / "09_raw_vs_cleaned_dataset_sizes.png",
        "Raw and cleaned dataset counts. Text and images retain all records in the cleaning summary; audio removes one invalid sample.",
    )
    add_heading(doc, "2.2 Cleaning as design quality control", 2)
    add_para(
        doc,
        "Cleaning is not an administrative step. Because this project uses data as design input, data errors can become visible design errors. A broken image path can produce a missing texture in Blender, an empty text fragment can weaken semantic retrieval, and an invalid audio duration can distort animation parameters. The notebooks therefore check whitespace, word count, short fragments, file existence, image validity, resolution, aspect ratio and duration.",
    )
    add_table_from_rows(
        doc,
        ["Dataset", "Quality check", "Result"],
        [
            ["Text", "Whitespace normalisation, word-count validation, short-fragment check", "288 valid text segments retained"],
            ["Images", "File existence, image validity, resolution and aspect-ratio validation", "298 valid image records retained"],
            ["Audio", "File existence and duration-based validation", "270 valid audio samples retained"],
        ],
        widths=[2.5, 8.2, 5.0],
        caption="Table 2. Cleaning and quality-control summary.",
    )
    add_two_images(
        doc,
        FIG / "part1_data_processing" / "02_cleaned_text_segments_per_chapter.png",
        "Cleaned text segments per chapter.",
        FIG / "part1_data_processing" / "08_cleaned_audio_duration_distribution.png",
        "Cleaned audio duration distribution.",
    )
    add_para(
        doc,
        "The cleaned datasets preserve traceability. A generated fragment can be traced to a row in parameter_table.csv, then to cleaned features and finally to source material. This traceability is what separates the work from a manually assembled visual collage.",
    )
    add_page_break(doc)

    # Chapter 3
    add_heading(doc, "3. Text Vectorisation: Motifs and Semantic Retrieval", 1)
    add_heading(doc, "3.1 TF-IDF as an explicit motif map", 2)
    add_para(
        doc,
        "The text dataset was vectorised using TF-IDF and all-MiniLM-L6-v2 sentence embeddings. TF-IDF generates a 300-dimensional keyword representation. It is useful because it is interpretable: it reveals the explicit motifs that dominate the Wonderland corpus. The strongest terms include Alice, Queen, Time, King, Gryphon, Turtle, Hatter, Rabbit, Head, Duchess, Mouse, Dormouse, Cat, Door, Caterpillar and White.",
    )
    add_image(
        doc,
        FIG / "part1_vectorisation" / "01_text_top_tfidf_terms.png",
        "Top TF-IDF terms. The keyword layer identifies explicit Wonderland motifs and characters.",
    )
    add_para(
        doc,
        "However, TF-IDF is limited because it treats similarity mainly through shared vocabulary. This becomes a problem for Wonderland because important themes are often indirect. A passage about dream logic may not repeatedly say dream, and a passage about time may involve the Hatter and tea-party repetition rather than only the word time.",
    )
    add_heading(doc, "3.2 MiniLM as semantic narrative similarity", 2)
    add_para(
        doc,
        "MiniLM sentence embeddings generate a 384-dimensional semantic representation. This model is less transparent than TF-IDF, but it is more useful for narrative similarity because it can relate fragments that express similar ideas through different words. In the current comparison, MiniLM produces a higher silhouette score than TF-IDF: 0.038866 compared with -0.008620.",
    )
    add_table_from_rows(
        doc,
        ["Method", "Dimensions", "Silhouette", "Interpretation"],
        [
            ["TF-IDF", "300", "-0.008620", "Readable keywords, but weak cluster separation"],
            ["all-MiniLM-L6-v2", "384", "0.038866", "Better semantic grouping, still overlapping"],
        ],
        widths=[4.0, 2.6, 2.6, 6.8],
        caption="Table 3. Text model comparison.",
    )
    add_image(
        doc,
        FIG / "part1_vectorisation" / "02_text_vectorisation_pca_comparison.png",
        "Text PCA comparison. MiniLM captures more structure in the first two components than TF-IDF, but both remain diffuse.",
    )
    add_heading(doc, "3.3 Retrieval evidence: why semantic search matters", 2)
    add_para(
        doc,
        "The retrieval results make the difference between keyword and semantic models especially clear. For the query mad tea party, TF-IDF ranks a generic passage containing party from The Pool of Tears as the top result. MiniLM retrieves the actual A Mad Tea-Party chapter as the top result. For the query dream, TF-IDF produces zero-score results, while MiniLM retrieves the final chapter where Alice wakes and describes the experience as a curious dream. These are not small technical details; they show that semantic embeddings are more aligned with Wonderland's dream logic.",
    )
    retrieval_rows = [
        ["mad tea party", "TF-IDF", "Generic party reference from The Pool of Tears", "Keyword confusion"],
        ["mad tea party", "MiniLM", "A Mad Tea-Party chapter retrieved first", "Scene-level meaning"],
        ["dream", "TF-IDF", "Zero-score unrelated passages", "Keyword failure"],
        ["dream", "MiniLM", "Final wake-up / curious dream passage", "Semantic success"],
        ["door", "Both", "Locked-door and threshold passages", "Strong explicit motif"],
        ["mushroom", "Both", "Caterpillar / mushroom scale-change scenes", "Strong object motif"],
    ]
    add_table_from_rows(
        doc,
        ["Query", "Model", "Result pattern", "Report interpretation"],
        retrieval_rows,
        widths=[3.0, 2.4, 6.1, 4.5],
        caption="Table 4. Retrieval comparison for Wonderland motifs.",
    )
    add_callout(
        doc,
        "Text analysis conclusion",
        "TF-IDF should be used as a readable motif layer. MiniLM should be used for semantic scene retrieval and narrative similarity. Low text cluster separation should be interpreted as part of Wonderland's ambiguous dream structure, not simply as a failed model.",
    )
    add_page_break(doc)

    # Chapter 4
    add_heading(doc, "4. Image and Audio Vectorisation: Atmosphere and Motion", 1)
    add_heading(doc, "4.1 Image features as visual atmosphere", 2)
    add_para(
        doc,
        "The image dataset was represented through 35 handcrafted colour and size features. The image clustering silhouette score is 0.170996, which is stronger than the text scores. This suggests that visual descriptors form clearer groups than narrative fragments. In design terms, image features are useful for atmosphere and material logic rather than symbolic recognition.",
    )
    add_two_images(
        doc,
        FIG / "part1_data_processing" / "06_curated_wonderland_image_showcase.png",
        "Curated image samples showing visual atmosphere.",
        FIG / "part1_vectorisation" / "04_image_handcrafted_pca_projection.png",
        "Handcrafted image-feature PCA projection.",
    )
    add_para(
        doc,
        "This distinction is important. Handcrafted features can describe colour, contrast, brightness, size and aspect ratio, but they do not automatically know whether an image contains a rabbit, teacup, mushroom or playing card. Therefore, image features should guide material, roughness, colour intensity and density, while text retrieval remains responsible for symbolic meaning.",
    )
    add_image(
        doc,
        FIG / "part1_vectorisation" / "04_image_feature_correlation.png",
        "Image feature correlation. Visual relationships can inform material and atmosphere mapping.",
    )
    add_heading(doc, "4.2 Audio features as motion and instability", 2)
    add_para(
        doc,
        "The audio dataset was represented through 11 basic descriptors and 26 MFCC features. Basic descriptors produce a stronger silhouette score than MFCCs: 0.199004 compared with 0.107957. The PCA variance comparison also supports this decision: the first two components of the basic descriptor model explain 39.99% and 28.26% of variance, while MFCC components explain 30.80% and 16.98%.",
    )
    add_table_from_rows(
        doc,
        ["Audio model", "Dimensions", "Silhouette", "PC1 + PC2 variance", "Design use"],
        [
            ["Basic descriptors", "11", "0.199004", "68.24%", "Motion amplitude, speed, direct TD parameters"],
            ["MFCC features", "26", "0.107957", "47.78%", "Secondary sound texture and timbral grouping"],
        ],
        widths=[3.5, 2.2, 2.4, 3.0, 5.0],
        caption="Table 5. Audio feature comparison.",
    )
    add_image(
        doc,
        FIG / "part1_vectorisation" / "08_audio_vectorisation_pca_comparison.png",
        "Audio PCA comparison. Basic descriptors form a more directly usable motion-control layer than MFCC features.",
    )
    add_para(
        doc,
        "Within the Alice/Wonderland theme, audio is interpreted as instability and rhythm: the White Rabbit's urgency, the Mad Tea-Party's broken time, the Queen's commands and Alice's sudden changes in scale. For this reason, basic audio descriptors become the primary source for motion_amp, anim_speed and temporal_sharpness.",
    )
    add_page_break(doc)

    # Chapter 5
    add_heading(doc, "5. Cross-Modal Model Comparison and Design Decision", 1)
    add_para(
        doc,
        "The model comparison shows that each modality performs a different design task. Text is semantically rich but difficult to cluster cleanly. Images cluster more clearly through visual descriptors and are useful for atmosphere and material. Audio descriptors produce the strongest clustering result and are most useful for motion mapping. The project therefore should not search for a single winning model across all data. It should choose the representation that best supports each design decision.",
    )
    add_table_from_rows(
        doc,
        ["Modality", "Preferred representation", "Reason", "Design decision"],
        [
            ["Text", "MiniLM + TF-IDF", "Semantic retrieval improves scene logic; TF-IDF remains readable", "Meaning, motifs, emotion, prompts"],
            ["Images", "Handcrafted colour/size", "Clearer visual grouping and direct material mapping", "Atmosphere, colour, roughness, texture"],
            ["Audio", "Basic descriptors", "Highest silhouette and clearer PCA structure", "Motion amplitude, speed, temporal behaviour"],
            ["MFCC", "Secondary only", "Useful timbral texture but less direct for animation", "Optional sound texture grouping"],
        ],
        widths=[2.8, 3.7, 5.2, 4.4],
        caption="Table 6. Cross-modal design decision.",
    )
    add_image(
        doc,
        GEN / "method_comparison_matrix.png",
        "Method comparison matrix. Models are evaluated through score, interpretability and design usefulness.",
    )
    add_callout(
        doc,
        "Case-level interpretation",
        "The low text score is not hidden because it explains the source material. Wonderland is ambiguous and overlapping. The stronger design strategy is not hard clustering, but soft semantic mapping into emotion, scene role and parameter tendencies.",
    )
    add_heading(doc, "5.1 Data-to-design mapping", 2)
    add_para(
        doc,
        "The output of Part 1 becomes operational through the parameter table. This table is the design interface of the project: it translates source records into geometry, material and animation controls. It is also the strongest proof that the notebooks are not isolated analysis exercises.",
    )
    param = pd.read_csv(pth("data", "processed", "parameters", "parameter_table.csv"))
    preview_cols = ["fragment_id", "scene_label", "emotion", "height", "complexity", "fragment_family", "motion_amp", "temporal_sharpness", "anim_speed"]
    preview_rows = param[preview_cols].head(6).round(3).astype(str).values.tolist()
    add_table_from_rows(
        doc,
        preview_cols,
        preview_rows,
        widths=[2.2, 2.4, 2.0, 1.5, 1.8, 1.8, 1.7, 2.2, 1.7],
        caption="Table 7. Excerpt from parameter_table.csv showing the bridge from Part 1 analysis to design parameters.",
    )
    add_image(
        doc,
        GEN / "parameter_distribution_grid.png",
        "Parameter distributions. Variation in height, complexity, motion and speed supports a diverse fragment library.",
    )
    add_part1_extended_figure_analysis(doc)
    add_page_break(doc)

    # Chapter 7
    add_heading(doc, "7. Workflow 1: Blender Fragment Library", 1)
    add_para(
        doc,
        "The Blender workflow converts the parameter table into a scripted fragment library. The script reads parameter_table.csv and texture_mapping.csv, clears the scene, builds lighting, sets a camera path and instantiates parameter rows as animated fragments. The workflow is repeatable: changing the parameter table changes the generated spatial output.",
    )
    add_image(
        doc,
        GEN / "blender_script_workflow_diagram.png",
        "Blender script workflow. CSV parameters are translated into primitive type, scale, material and animation.",
    )
    add_para(
        doc,
        "The mapping inside Blender is direct. fragment_family selects the primitive type. width, depth and height define object scale. complexity controls subdivision. roughness affects material response. emotion provides fallback colour when texture paths are unavailable. scene_order_param changes circular radius and vertical placement. motion_amp drives the floating animation after growth.",
    )
    add_two_images(
        doc,
        STILLS / "06_part2_workflow1_blender_fragments" / "stills" / "wf1_still_01.png",
        "Blender workflow still: early or establishing view.",
        STILLS / "06_part2_workflow1_blender_fragments" / "stills" / "wf1_still_03.png",
        "Blender workflow still: fragment variation and composition.",
    )
    add_callout(
        doc,
        "Blender analysis",
        "The output should be read as a visualisation of mapped data. Primitive variation represents family classification, scale represents geometric parameters, material variation represents image and emotion inputs, and floating movement represents temporal energy.",
    )

    # Chapter 8
    add_heading(doc, "8. Workflow 2: TouchDesigner Procedural System", 1)
    add_para(
        doc,
        "TouchDesigner provides a different design environment from Blender. Instead of a scripted cinematic render, it supports live node-based manipulation. This makes it useful for testing procedural behaviours, parameter remapping and point-cloud or instancing systems. Including TouchDesigner strengthens the project because it demonstrates that the parameter table is not tied to one software package.",
    )
    add_two_images(
        doc,
        STILLS / "TD capture" / "1.png",
        "TouchDesigner network capture: data import or preparation.",
        STILLS / "TD capture" / "3.png",
        "TouchDesigner network capture: instancing or geometry logic.",
    )
    add_two_images(
        doc,
        STILLS / "07_part2_workflow2_touchdesigner_pointcloud" / "stills" / "wf2_still_01.png",
        "TouchDesigner output still: procedural visual state.",
        STILLS / "07_part2_workflow2_touchdesigner_pointcloud" / "stills" / "wf2_still_03.png",
        "TouchDesigner output still: later procedural state.",
    )
    add_para(
        doc,
        "The difference between Blender and TouchDesigner is methodological. Blender provides a controlled rendered sequence with lighting and camera movement. TouchDesigner provides a live procedural system where values can be remapped and tested visually. Together, they show two interpretations of the same parameter system: cinematic and procedural.",
    )
    add_page_break(doc)

    # Chapter 9
    add_heading(doc, "9. Final Animation and Critical Reflection", 1)
    add_heading(doc, "9.1 Combined output", 2)
    add_para(
        doc,
        "The final animation combines the project workflows into a spatial narrative. The opening establishes atmosphere, the middle section shows data-driven variation through fragments and motion, and the final section synthesises the outputs into a complete visual statement. The animation should be read as the proof that the Part 1 data work can drive Part 2 design behaviour.",
    )
    add_two_images(
        doc,
        STILLS / "09_final_animation_reflection" / "stills" / "final_still_01.png",
        "Final animation still: opening or establishing state.",
        STILLS / "09_final_animation_reflection" / "stills" / "final_still_03.png",
        "Final animation still: later transformation or synthesis.",
    )
    add_heading(doc, "9.2 Strengths", 2)
    add_para(
        doc,
        "The strongest part of the project is the explicit bridge from cleaned data to spatial output. The notebooks produce more than charts: they produce design instructions. The model comparison is also valuable because it explains why certain feature sets are used. This gives the final animation a methodological basis rather than only an aesthetic direction.",
    )
    add_heading(doc, "9.3 Limitations", 2)
    add_para(
        doc,
        "The main limitations are specific. Text clustering scores are low, which suggests overlapping narrative meanings. Image analysis currently relies on handcrafted descriptors rather than semantic image embeddings such as CLIP. Texture mapping can break if absolute file paths are moved between machines. The Blender script currently uses a subset of parameter records rather than the full table. These limitations should be stated clearly because they show technical awareness.",
    )
    add_heading(doc, "9.4 Future development", 2)
    add_para(
        doc,
        "Future work could add CLIP image embeddings for semantic object recognition, improve audio segmentation, render all 50 parameter records, store texture paths relatively and introduce richer deformation driven by text and image features. These improvements would deepen the link between source data and spatial behaviour while preserving the current workflow structure.",
    )
    add_callout(
        doc,
        "Reflection conclusion",
        "The project is best understood as data-informed authorship. The data does not automatically create Wonderland; it structures the designer's choices and makes the final spatial narrative traceable.",
    )
    add_page_break(doc)

    # Appendix
    add_heading(doc, "Appendix: Evidence Checklist and Credits", 1)
    add_table_from_rows(
        doc,
        ["Evidence", "Location", "Status / note"],
        [
            ["Notebook 01", "notebooks/PART1/01_data_collection_enhanced.ipynb", "Data collection and raw dataset evidence"],
            ["Notebook 02", "notebooks/PART1/02_data_processing.ipynb", "Cleaning, validation and processed datasets"],
            ["Notebook 03", "notebooks/PART1/03_vectorisation_and_comparison_enhanced_final.ipynb", "Vectorisation, PCA and model comparison"],
            ["Notebook 04", "notebooks/PART1/04_parameter_mapping.ipynb", "Parameter table and design mapping"],
            ["Parameter table", "data/processed/parameters/parameter_table.csv", "Main bridge between analysis and design production"],
            ["Blender workflow", "outputs/animation/video-wf1.mp4", "Scripted fragment library"],
            ["TouchDesigner workflow", "outputs/animation/video-wf2.mp4", "Procedural point-cloud / node system"],
            ["Final animation", "outputs/animation/3D_animation_combined.mp4", "Combined final output"],
        ],
        widths=[3.4, 8.2, 5.0],
        caption="Table 9. Submission evidence checklist.",
    )
    add_heading(doc, "Credits and tools", 2)
    add_para(
        doc,
        "Technical stack: Python, pandas, scikit-learn, sentence-transformers, image/audio feature extraction, matplotlib/seaborn outputs from the notebooks, Blender Python API, TouchDesigner and video export tools. Source data, API/model use, AI-assisted tools and final repository/OneDrive links should be credited in the final submission package.",
    )
    add_heading(doc, "Final pre-submission check", 2)
    add_para(
        doc,
        "Before submission, check that every figure has a caption, every dataset count is consistent, all notebook paths are correct, texture paths are portable or explained, the final animation exceeds the required duration, and GitHub/OneDrive links are accessible.",
    )

    # Header/footer
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("Reconstructing Wonderland | Data-Driven Spatial Narrative")
        style_run(run, size=8.5, italic=True, color="666666")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
